"""
document_processor.py — ⽂档加载 & 语义分块
⽀持： PDF / DoCX / TXT / MD / 纯⽂本字符串
"""
import os
import json
import hashlib
from typing import List, Optional
from dataclasses import dataclass, field

from langchain_text_splitters import RecursiveCharacterTextSplitter
from rich.console import Console

from rag_config import RAGConfig, DEFAULT_CONFIG

console = Console()

# ─────────────────────────────────────────────
# 数据结构
# ─────────────────────────────────────────────
@dataclass
class Document:
    """单个⽂档块"""
    doc_id: str          # 唯⼀ ID（ 内容哈希）
    content: str         # ⽂本内容
    source: str          # 来源⽂件路径
    chunk_index: int     # 在原⽂中的块序号
    metadata: dict = field(default_factory=dict)

    def to_dict(self) -> dict:
        return {
            "doc_id": self.doc_id,
            "content": self.content,
            "source": self.source,
            "chunk_index": self.chunk_index,
            "metadata": self.metadata,
        }

    @staticmethod
    def from_dict(d: dict) -> "Document":
        return Document(**d)

# ─────────────────────────────────────────────
# ⽂档加载器
# ─────────────────────────────────────────────
class DocumentLoader:
    """从⽂件或字符串加载原始⽂本"""
    @staticmethod
    def load_file(path: str) -> str:
        ext = os.path.splitext(path)[-1].lower()

        if ext == ".txt" or ext == ".md":
            with open(path, "r", encoding="utf-8") as f:
                return f.read()

        elif ext == ".pdf":
            from pypdf import PdfReader
            reader = PdfReader(path)
            return "\n".join(page.extract_text() or "" for page in reader.pages)

        elif ext == ".docx":
            from docx import Document as DocxDocument
            doc = DocxDocument(path)
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

        else:
            raise ValueError(f"不⽀持的⽂件类型: {ext}， 请使⽤ .txt/.md/.pdf/.docx")

    @staticmethod
    def load_directory(dir_path: str) -> List[tuple]:
        """加载⽬录下所有⽀持的⽂档， 返回 [(path, text), ...]"""
        supported = {".txt", ".md", ".pdf", ".docx"}
        results = []
        for fname in os.listdir(dir_path):
            if os.path.splitext(fname)[-1].lower() in supported:
                fpath = os.path.join(dir_path, fname)
                try:
                    text = DocumentLoader.load_file(fpath)
                    results.append((fpath, text))
                    console.print(f"[green]✓ 加载[/green] {fname} ({len(text)} 字符)")
                except Exception as e:
                    console.print(f"[red]✗ 跳过[/red] {fname}: {e}")
        return results

# ─────────────────────────────────────────────
# 分块处理器
# ─────────────────────────────────────────────
class DocumentProcessor:
    """
    ⽂档预处理主类
    流程： 加载⽂本 → 清洗 → 语义分块 → ⽣成 Document 列表
    """
    def __init__(self, config: RAGConfig = DEFAULT_CONFIG):
        self.cfg = config.chunk
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.cfg.chunk_size,
            chunk_overlap=self.cfg.chunk_overlap,
            separators=self.cfg.separators,
            length_function=self._token_len,
        )

    @staticmethod
    def _token_len(text: str) -> int:
        """⽤ tiktoken 估算 token 数， 中⽂场景下⽤字符数 / 1.5 粗估"""
        try:
            import tiktoken
            enc = tiktoken.get_encoding("cl100k_base")
            return len(enc.encode(text))
        except Exception:
            return len(text) // 2  # fallback 兜底方案

    @staticmethod
    def _clean(text: str) -> str:
        """基础⽂本清洗"""
        import re
        # 删除多余空⽩⾏
        text = re.sub(r"\n{3,}", "\n\n", text)
        # 删除⻚眉⻚脚常⻅符号
        text = re.sub(r"\.{4,}", "", text)
        # 删除零宽字符
        text = re.sub(r"[\u200b\u200c\u200d\ufeff]", "", text)
        # 去除首尾空白
        return text.strip()

    @staticmethod
    def _make_id(source: str, index: int, content: str) -> str:
        """⽣成稳定唯⼀ ID"""
        raw = f"{source}::{index}::{content[:50]}"
        return hashlib.md5(raw.encode()).hexdigest()

    def process_text(self, text: str, source: str = "inline") -> List[Document]:
        """处理纯⽂本字符串"""
        text = self._clean(text)
        chunks = self.splitter.split_text(text)
        docs = []

        for i, chunk in enumerate(chunks):
            if len(chunk.strip()) < 10:  # 过滤太短的 chunk，避免无效内容
                continue
            doc = Document(
                doc_id=self._make_id(source, i, chunk),
                content=chunk,
                source=source,
                chunk_index=i,
                metadata={"total_chunks": len(chunks)},
            )
            docs.append(doc)
        return docs

    def process_file(self, path: str) -> List[Document]:
        """处理单个⽂件"""
        text = DocumentLoader.load_file(path)
        return self.process_text(text, source=path)

    def process_directory(self, dir_path: str) -> List[Document]:
        """处理整个⽬录"""
        all_docs = []
        file_pairs = DocumentLoader.load_directory(dir_path)
        for path, text in file_pairs:
            docs = self.process_text(text, source=path)
            all_docs.extend(docs)
            console.print(f" → 切分为 [cyan]{len(docs)}[/cyan] 个 chunks")

        console.print(f"\n[bold green]共处理 {len(all_docs)} 个⽂档块[/bold green]")
        return all_docs

    @staticmethod
    def save_docs(docs: List[Document], path: str):
        """持久化⽂档块到 JSON"""
        os.makedirs(os.path.dirname(path), exist_ok=True)
        with open(path, "w", encoding="utf-8") as f:
            json.dump(
                [d.to_dict() for d in docs],
                f,
                ensure_ascii=False,
                indent=2
            )
        console.print(f"[green]⽂档块已保存到 {path}[/green]")

    @staticmethod
    def load_docs(path: str) -> List[Document]:
        """从 JSON 加载⽂档块"""
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return [Document.from_dict(d) for d in data]